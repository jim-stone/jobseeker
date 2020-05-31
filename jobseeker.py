#!/usr/bin/env python3

import requests
from docx import Document
from lxml import etree


"""
Wyszukuje ogłoszenia o odpowiednich stanowiskach pracy.
Znajduje informacje w pliku xml, po czym kopiuje 
do dokumentu word, nadając podstawowe formatowanie.
"""


URL = 'https://nabory.kprm.gov.pl/pls/serwis/app.xml'
RELEVANT_POSITIONS = ['naczelnik', 'kierownik', 'główny specjalista']
RELEVANT_FIELDS = ['nazwa_firmy', 'stanowisko', 'do_spraw', 'komorka_organizacyjna',
                'etykieta1', 'poledodatkowelista1', 'etykieta2',
                'pozostale_wym_niezbedne', 'etykieta_pozostale_wym_pozadane',
                'wym_pozadane', 'poledodatkowetext3'] 
TEXT_TO_FIND = 'new/oferta/stanowisko'
REPLACEMENTS = [('None', ''), ('<br>', ''), ('<br />', ''), ('<ul><li>', ''),
                ('</li><li>', '\n'), ('</li></ul>', '\n'), ('<ul></ul>', '')]


def main():
    root = get_xml_root (URL)
    document = Document ()
    try:
        copy_from_root_to_document(XMLroot=root, document=document)
    except Exception:
        print("Wystąpił błąd. Nie udało się pobrać pliku.")
    else:
        print("Plik z ofertami pracy został pobrany.")


def get_xml_root (url=URL):
    try:
        response = requests.get (url)
        if response.status_code != 200:
            print (response.status_code)
        else:
            content = response.content
            root = etree.XML (content)
            return root
    except Exception as e:
        print (e)


def copy_from_root_to_document (XMLroot, document):
    
    for element in XMLroot.findall (TEXT_TO_FIND):
        position = element.text.split (' ', 1)[1][:-4]
        if position in RELEVANT_POSITIONS:
            for sib in element.getparent().getchildren():               
                pole = sib.tag
                description = str (sib.text)

                if pole in RELEVANT_FIELDS [0:4]:
                    if 'CDATA' in description:
                        description = split_description(description)
                    description = replace_in_description(description)
                    para = document.add_paragraph (description)
                    try:
                        para.runs[0].bold = True
                    except:
                        pass
                
                elif pole in RELEVANT_FIELDS [-1]:
                    description = split_description(description)
                    para = document.add_paragraph (description)
                    para.runs[0].bold=True
                    document.add_page_break ()
                
                elif pole in RELEVANT_FIELDS:
                    if 'CDATA' in description:
                        description = split_description(description)
                    description = replace_in_description(description)
                    document.add_paragraph (description)

    document.save ('kprm_oferty.doc')


def split_description(description):
    return description.split (' ', 1)[1][:-4]

def replace_in_description(text):
    for pair in REPLACEMENTS:
        text = text.replace(pair[0], pair[1])
    return text


if __name__ == '__main__':
    main()